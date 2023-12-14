from docx import Document
from docx.shared import Pt, RGBColor
import string


class PostProcessor:
    def __init__(self,
                 template_docx_path):
        self.document = Document(template_docx_path)

    def remove_punctuation(self, text):
        # Remove punctuation from the text
        translator = str.maketrans("", "", string.punctuation)
        return text.translate(translator)


    def write_doc(self,
                  responses,output_path):
        for paragraph in self.document.paragraphs:
            if paragraph.text == "candidate_name":
                paragraph.text = self.remove_punctuation(responses["candidate_name"].strip("The name of the candidate mentioned in the given text is "))
                for run in paragraph.runs:
                    run.font.size = Pt(22)  # Set font size to 22 points
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black

            if paragraph.text == "candidate_title":
                paragraph.text = self.remove_punctuation(responses["candidate_title"].strip("The candidate job title mentioned in the given text is "))
                for run in paragraph.runs:
                    run.font.size = Pt(20)  # Set font size to 20 points
                    run.font.color.rgb = RGBColor(255, 0,
                                                  0)  # Set font color to red (255 for red, 0 for green, 0 for blue)
                    run.font.bold = None  # Remove bold formatting

        for i, table in enumerate(self.document.tables):
            if i == 0:
                table.cell(0, 1).add_paragraph(responses["career_summary"])
            if i == 1:
                table.cell(0, 1).add_paragraph(responses["skills"].split("Based on the given resume, the top 10 technical skills listed are:")[-1])
            if i == 2:
                table.cell(0,1).add_paragraph(responses["work_experience"])

            if i==3:
                table.cell(0,1).add_paragraph(responses["projects"])

            if i == 4:
                degrees = responses["education"].split("\n")
                for degree in degrees:
                    table.cell(0,1).add_paragraph(degree)

        self.document.save(output_path)