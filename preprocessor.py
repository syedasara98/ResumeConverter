from docx import Document

class Preprocessor:
    def __init__(self):
        pass

    @staticmethod
    def convert_docx_to_text(docx_path):
        # Open the Word document
        doc = Document(docx_path)
        texts = []
        text = ""
        for paragraph in doc.paragraphs:
            # Write each paragraph to the text file
            print(paragraph.text)
            texts.append(paragraph.text + "\n")
        text = "".join(texts)
        return text

    # Create or open a text file for writing
    # with open(txt_path, 'w', encoding='utf-8') as txt_file:
    #     # Iterate through paragraphs in the Word document
    #     for paragraph in doc.paragraphs:
    #         # Write each paragraph to the text file
    #         print(paragraph.text)
    #
    #         # txt_file.write(paragraph.text + '\n')
#
#     print(f"Conversion completed. Text saved to '{txt_path}'")
#
# # Replace 'input_document.docx' with the path to your Word document
# input_docx_path = '/home/codeninja/Hackathon/Sample Resumes/Muhammad Amjad Resume Updated.docx'
# input_docx_path = '../Sample Resumes/Waseem-Haider.docx'
#
# # Replace 'output_text.txt' with the desired path for the plain text file
# output_txt_path = 'output_text.txt'
#
# # Call the function to convert Word document to plain text
# convert_docx_to_text(input_docx_path, output_txt_path)
